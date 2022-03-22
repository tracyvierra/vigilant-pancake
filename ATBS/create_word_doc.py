# Author: Tracy Vierra
# Date Created: 3/21/2022
# Date Modified: 3/21/2022

# Description:

# Usage:

import docx
import os

os.chdir('C:\\Users\\tracy\\Documents\\Python Scripts\\git\\vigilant-pancake\\ATBS')

d = docx.Document()
d.add_paragraph('Hello world!')
d.add_paragraph('This is a second paragraph.')
d.add_paragraph('This is a third paragraph.')
d.add_paragraph('This is a fourth paragraph.')
d.add_paragraph('This is a fifth paragraph.')
d.add_paragraph('This is a sixth paragraph.')

d.save('demo4.docx')
p = d.paragraphs[0]
p.add_run(' This is a new run.')
p.runs[1].add_break()
p.runs[1].add_break()
p.runs[1].bold = True

d.save('demo5.docx')

def get_text(filename):
	doc = docx.Document(filename)
	full_text = []
	for para in doc.paragraphs:
		full_text.append(para.text)
		return '\n'.join(full_text)


print(get_text('demo5.docx'))
print(get_text('demo.docx'))

