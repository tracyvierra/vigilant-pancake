# Author: Tracy Vierra
# Date Created: 3/21/2022
# Date Modified: 3/21/2022

# Description:

# Usage:

import docx
import os

os.chdir('C:\\Users\\tracy\\Documents\\Python Scripts\\git\\vigilant-pancake\\ATBS')

d = docx.Document('demo.docx')
print(d.paragraphs)
print(d.paragraphs[0].text)
print(d.paragraphs[1].text)
print(d.paragraphs[2].text)
print(d.paragraphs[3].text)
print(d.paragraphs[4].text)
print(d.paragraphs[5].text)
print(d.paragraphs[6].text)

p = d.paragraphs[1]
print(p.runs)
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)

print(p.runs[0].bold)
print(p.runs[0].italic)
print(p.runs[0].underline)
print(p.runs[0].font.name)
print(p.runs[0].font.size)
print(p.runs[0].font.color.rgb)

print(p.runs[1].bold)
print(p.runs[1].italic)
print(p.runs[1].underline)
print(p.runs[1].font.name)
print(p.runs[1].font.size)
print(p.runs[1].font.color.rgb)

print(p.runs[3].bold)
print(p.runs[3].italic)
print(p.runs[3].underline)
print(p.runs[3].font.name)
print(p.runs[3].font.size)
print(p.runs[3].font.color.rgb)

p.runs[3].text = 'italic and underline'
p.runs[3].underline = True
p.runs[3].italic = True

d.save('demo_edited.docx')



print(p.style)

